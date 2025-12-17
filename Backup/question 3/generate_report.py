"""
Script to generate a professional Word document report from the News Classification project
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_report():
    """Create a comprehensive professional report"""
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ==================== TITLE PAGE ====================
    title = doc.add_heading('News Classification Project', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('ML/DL with Reinforcement Learning', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Student Information
    p = doc.add_paragraph()
    p.add_run('Name: ').bold = True
    p.add_run('ARINEITWE THOMAS')
    
    p = doc.add_paragraph()
    p.add_run('Lecturer: ').bold = True
    p.add_run('Dr Sibitenda')
    
    doc.add_page_break()
    
    # ==================== TABLE OF CONTENTS ====================
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. Executive Summary',
        '2. Project Overview',
        '3. Part A: Data Mining and Pre-processing',
        '4. Part B: Two News Classifiers',
        '5. Part C: Topic Clustering (TF-IDF)',
        '6. Part D: Reinforcement Learning Decision Agent',
        '7. Results and Analysis',
        '8. Conclusion',
        '9. References'
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # ==================== EXECUTIVE SUMMARY ====================
    doc.add_heading('1. Executive Summary', 1)
    doc.add_paragraph(
        'This project implements a comprehensive news article classification system that combines '
        'classical machine learning, deep learning, and reinforcement learning techniques. The system '
        'successfully classifies BBC news articles into five categories (business, entertainment, '
        'politics, sport, tech), discovers latent topics through clustering, and uses a reinforcement '
        'learning agent to intelligently decide which model to use or when to escalate to human review.'
    )
    
    doc.add_paragraph('Key Achievements:', style='List Bullet')
    doc.add_paragraph('• Classical ML Model (Logistic Regression): 98.88% accuracy', style='List Bullet 2')
    doc.add_paragraph('• Deep Learning Model (CNN): 95.06% accuracy', style='List Bullet 2')
    doc.add_paragraph('• Reinforcement Learning Agent: 99.10% accuracy', style='List Bullet 2')
    doc.add_paragraph('• Successfully identified 5 distinct topic clusters', style='List Bullet 2')
    
    doc.add_page_break()
    
    # ==================== PROJECT OVERVIEW ====================
    doc.add_heading('2. Project Overview', 1)
    
    doc.add_heading('2.1 Objectives', 2)
    doc.add_paragraph(
        'The primary objectives of this project are:'
    )
    objectives = [
        'Classify news articles into 5 categories using classical machine learning',
        'Implement a deep learning model for news classification',
        'Discover latent topics through clustering analysis',
        'Develop a reinforcement learning agent to intelligently select the best model or escalate to human review'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    doc.add_heading('2.2 Dataset Overview', 2)
    doc.add_paragraph(
        'Dataset: BBC News Articles'
    )
    doc.add_paragraph(
        'Source: https://www.kaggle.com/datasets/pariza/bbc-news-summary'
    )
    
    # Dataset Statistics Table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    stats = [
        ('Total Articles', '2,225'),
        ('Categories', '5 (business, entertainment, politics, sport, tech)'),
        ('Format', 'Text files organized in category folders'),
        ('Distribution', 'Balanced across categories')
    ]
    
    for i, (label, value) in enumerate(stats):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_paragraph(
        '[INSERT SCREENSHOT: Class Distribution Bar Chart - Figure showing distribution of articles across 5 categories]'
    )
    
    # ==================== PART A: DATA MINING AND PRE-PROCESSING ====================
    doc.add_page_break()
    doc.add_heading('3. Part A: Data Mining and Pre-processing', 1)
    
    doc.add_heading('3.1 Data Loading', 2)
    doc.add_paragraph(
        'The dataset was loaded from the local directory structure where articles are organized '
        'by category folders. A total of 2,225 articles were successfully loaded across 5 categories.'
    )
    
    doc.add_paragraph('Dataset Statistics:', style='List Bullet')
    doc.add_paragraph('• Total articles: 2,225', style='List Bullet 2')
    doc.add_paragraph('• Categories: 5', style='List Bullet 2')
    doc.add_paragraph('• Category distribution:', style='List Bullet 2')
    
    category_dist = [
        'Business: 510 articles',
        'Entertainment: 386 articles',
        'Politics: 417 articles',
        'Sport: 511 articles',
        'Tech: 401 articles'
    ]
    for dist in category_dist:
        doc.add_paragraph(dist, style='List Bullet 3')
    
    doc.add_paragraph()
    doc.add_paragraph(
        '[INSERT SCREENSHOT: Sample Articles Display - Showing 5 sample articles with their labels]'
    )
    
    doc.add_heading('3.2 Text Pre-processing', 2)
    doc.add_paragraph(
        'Text cleaning was performed to prepare the data for analysis. The cleaning process included:'
    )
    
    cleaning_steps = [
        'Converting text to lowercase',
        'Removing HTML tags',
        'Removing digits and punctuation',
        'Removing extra whitespace'
    ]
    for step in cleaning_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Example:', style='List Bullet')
    doc.add_paragraph('Original: "Ad sales boost Time Warner profit\nQuarterly profits at US media giant TimeWarner jumped 76% to $1.1..."', style='List Bullet 2')
    doc.add_paragraph('Cleaned: "ad sales boost time warner profit quarterly profits at us media giant timewarner jumped to bn..."', style='List Bullet 2')
    
    doc.add_heading('3.3 Feature Engineering', 2)
    
    doc.add_heading('3.3.1 TF-IDF Features', 3)
    doc.add_paragraph(
        'TF-IDF (Term Frequency-Inverse Document Frequency) vectorization was used to create '
        'features for the classical machine learning model.'
    )
    
    tfidf_table = doc.add_table(rows=3, cols=2)
    tfidf_table.style = 'Light Grid Accent 1'
    tfidf_params = [
        ('Max Features', '5,000'),
        ('N-gram Range', '(1, 2)'),
        ('Stop Words', 'English')
    ]
    for i, (param, value) in enumerate(tfidf_params):
        tfidf_table.rows[i].cells[0].text = param
        tfidf_table.rows[i].cells[1].text = value
        tfidf_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph('Result: TF-IDF matrix shape: (2,225 × 5,000)')
    
    doc.add_heading('3.3.2 Tokenized Sequences for Deep Learning', 3)
    doc.add_paragraph(
        'For the deep learning model, text was tokenized and converted to sequences.'
    )
    
    dl_table = doc.add_table(rows=3, cols=2)
    dl_table.style = 'Light Grid Accent 1'
    dl_params = [
        ('Max Vocabulary Size', '10,000'),
        ('Max Sequence Length', '300'),
        ('Vocabulary Size', '31,519')
    ]
    for i, (param, value) in enumerate(dl_params):
        dl_table.rows[i].cells[0].text = param
        dl_table.rows[i].cells[1].text = value
        dl_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph('Result: Sequence tensor shape: (2,225 × 300)')
    
    doc.add_heading('3.4 Train/Test Split', 2)
    doc.add_paragraph(
        'The dataset was split into training and testing sets using stratified sampling to maintain '
        'class distribution.'
    )
    
    split_table = doc.add_table(rows=3, cols=3)
    split_table.style = 'Light Grid Accent 1'
    split_table.rows[0].cells[0].text = 'Dataset'
    split_table.rows[0].cells[1].text = 'Size'
    split_table.rows[0].cells[2].text = 'Percentage'
    split_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    split_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    split_table.rows[0].cells[2].paragraphs[0].runs[0].bold = True
    
    split_table.rows[1].cells[0].text = 'Training'
    split_table.rows[1].cells[1].text = '1,780'
    split_table.rows[1].cells[2].text = '80.0%'
    
    split_table.rows[2].cells[0].text = 'Testing'
    split_table.rows[2].cells[1].text = '445'
    split_table.rows[2].cells[2].text = '20.0%'
    
    # ==================== PART B: TWO NEWS CLASSIFIERS ====================
    doc.add_page_break()
    doc.add_heading('4. Part B: Two News Classifiers', 1)
    
    doc.add_heading('4.1 Classical Machine Learning Model (Logistic Regression)', 2)
    
    doc.add_paragraph(
        'A Logistic Regression model was trained using TF-IDF features. This classical machine '
        'learning approach provides a baseline for comparison with the deep learning model.'
    )
    
    doc.add_heading('4.1.1 Model Configuration', 3)
    config_table = doc.add_table(rows=2, cols=2)
    config_table.style = 'Light Grid Accent 1'
    config_table.rows[0].cells[0].text = 'Parameter'
    config_table.rows[0].cells[1].text = 'Value'
    config_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    config_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    
    config_table.rows[1].cells[0].text = 'Max Iterations'
    config_table.rows[1].cells[1].text = '2,000'
    
    doc.add_heading('4.1.2 Results', 3)
    
    # Classification Report Table
    doc.add_paragraph('Classification Report:', style='List Bullet')
    
    ml_results_table = doc.add_table(rows=7, cols=5)
    ml_results_table.style = 'Light Grid Accent 1'
    ml_results_table.rows[0].cells[0].text = 'Category'
    ml_results_table.rows[0].cells[1].text = 'Precision'
    ml_results_table.rows[0].cells[2].text = 'Recall'
    ml_results_table.rows[0].cells[3].text = 'F1-Score'
    ml_results_table.rows[0].cells[4].text = 'Support'
    
    for cell in ml_results_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
    
    ml_results = [
        ('Business', '1.00', '0.96', '0.98', '102'),
        ('Entertainment', '1.00', '1.00', '1.00', '77'),
        ('Politics', '0.99', '0.99', '0.99', '84'),
        ('Sport', '0.99', '1.00', '1.00', '102'),
        ('Tech', '0.96', '1.00', '0.98', '80'),
        ('Overall', '0.99', '0.99', '0.99', '445')
    ]
    
    for i, (cat, prec, rec, f1, sup) in enumerate(ml_results, 1):
        ml_results_table.rows[i].cells[0].text = cat
        ml_results_table.rows[i].cells[1].text = prec
        ml_results_table.rows[i].cells[2].text = rec
        ml_results_table.rows[i].cells[3].text = f1
        ml_results_table.rows[i].cells[4].text = sup
    
    doc.add_paragraph()
    doc.add_paragraph('Key Metrics:', style='List Bullet')
    doc.add_paragraph('• Overall Accuracy: 98.88%', style='List Bullet 2')
    doc.add_paragraph('• Macro F1-Score: 0.9890', style='List Bullet 2')
    doc.add_paragraph('• Weighted F1-Score: 0.9887', style='List Bullet 2')
    doc.add_paragraph('• Total Misclassified: 5 articles', style='List Bullet 2')
    
    doc.add_paragraph()
    doc.add_paragraph(
        '[INSERT SCREENSHOT: Confusion Matrix for ML Model - Heatmap showing classification results]'
    )
    
    doc.add_heading('4.1.3 Misclassified Examples', 3)
    doc.add_paragraph(
        'The model misclassified 5 articles out of 445 test samples. Examples of misclassifications '
        'included articles that had overlapping themes between categories, such as business articles '
        'with technology terminology being classified as tech, or business articles with political '
        'content being classified as politics.'
    )
    
    doc.add_heading('4.2 Deep Learning Model (CNN)', 2)
    
    doc.add_paragraph(
        'A Convolutional Neural Network (CNN) was implemented using TensorFlow/Keras for deep learning-based '
        'classification. The model uses word embeddings and convolutional layers to capture local patterns in text.'
    )
    
    doc.add_heading('4.2.1 Model Architecture', 3)
    
    architecture_table = doc.add_table(rows=8, cols=2)
    architecture_table.style = 'Light Grid Accent 1'
    architecture_table.rows[0].cells[0].text = 'Layer'
    architecture_table.rows[0].cells[1].text = 'Configuration'
    architecture_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    architecture_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    
    arch_layers = [
        ('Embedding', 'Input: 10,000, Output: 100'),
        ('Conv1D', 'Filters: 128, Kernel: 5, Activation: ReLU'),
        ('GlobalMaxPooling1D', 'Pooling layer'),
        ('Dropout', 'Rate: 0.5'),
        ('Dense', 'Units: 64, Activation: ReLU'),
        ('Dropout', 'Rate: 0.5'),
        ('Dense (Output)', 'Units: 5, Activation: Softmax')
    ]
    
    for i, (layer, config) in enumerate(arch_layers, 1):
        architecture_table.rows[i].cells[0].text = layer
        architecture_table.rows[i].cells[1].text = config
    
    doc.add_heading('4.2.2 Training Configuration', 3)
    
    train_table = doc.add_table(rows=5, cols=2)
    train_table.style = 'Light Grid Accent 1'
    train_table.rows[0].cells[0].text = 'Parameter'
    train_table.rows[0].cells[1].text = 'Value'
    train_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    train_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    
    train_params = [
        ('Epochs', '20'),
        ('Batch Size', '32'),
        ('Optimizer', 'Adam'),
        ('Loss Function', 'Categorical Crossentropy')
    ]
    
    for i, (param, value) in enumerate(train_params, 1):
        train_table.rows[i].cells[0].text = param
        train_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    doc.add_paragraph(
        '[INSERT SCREENSHOT: Training Curves - Accuracy and Loss plots showing training progress over 20 epochs]'
    )
    
    doc.add_heading('4.2.3 Results', 3)
    
    dl_results_table = doc.add_table(rows=7, cols=5)
    dl_results_table.style = 'Light Grid Accent 1'
    dl_results_table.rows[0].cells[0].text = 'Category'
    dl_results_table.rows[0].cells[1].text = 'Precision'
    dl_results_table.rows[0].cells[2].text = 'Recall'
    dl_results_table.rows[0].cells[3].text = 'F1-Score'
    dl_results_table.rows[0].cells[4].text = 'Support'
    
    for cell in dl_results_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
    
    dl_results = [
        ('Business', '0.93', '0.89', '0.91', '102'),
        ('Entertainment', '1.00', '0.99', '0.99', '77'),
        ('Politics', '0.93', '0.93', '0.93', '84'),
        ('Sport', '0.98', '0.99', '0.99', '102'),
        ('Tech', '0.92', '0.96', '0.94', '80'),
        ('Overall', '0.95', '0.95', '0.95', '445')
    ]
    
    for i, (cat, prec, rec, f1, sup) in enumerate(dl_results, 1):
        dl_results_table.rows[i].cells[0].text = cat
        dl_results_table.rows[i].cells[1].text = prec
        dl_results_table.rows[i].cells[2].text = rec
        dl_results_table.rows[i].cells[3].text = f1
        dl_results_table.rows[i].cells[4].text = sup
    
    doc.add_paragraph()
    doc.add_paragraph('Key Metrics:', style='List Bullet')
    doc.add_paragraph('• Overall Accuracy: 95.06%', style='List Bullet 2')
    doc.add_paragraph('• Macro F1-Score: 0.9513', style='List Bullet 2')
    doc.add_paragraph('• Weighted F1-Score: 0.9504', style='List Bullet 2')
    doc.add_paragraph('• Total Misclassified: 22 articles', style='List Bullet 2')
    
    doc.add_paragraph()
    doc.add_paragraph(
        '[INSERT SCREENSHOT: Confusion Matrix for DL Model - Heatmap showing classification results]'
    )
    
    doc.add_heading('4.3 Model Comparison', 2)
    
    comparison_table = doc.add_table(rows=4, cols=3)
    comparison_table.style = 'Light Grid Accent 1'
    comparison_table.rows[0].cells[0].text = 'Model'
    comparison_table.rows[0].cells[1].text = 'Accuracy'
    comparison_table.rows[0].cells[2].text = 'F1-Score (Macro)'
    comparison_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    comparison_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    comparison_table.rows[0].cells[2].paragraphs[0].runs[0].bold = True
    
    comparison_table.rows[1].cells[0].text = 'Logistic Regression (ML)'
    comparison_table.rows[1].cells[1].text = '98.88%'
    comparison_table.rows[1].cells[2].text = '0.9890'
    
    comparison_table.rows[2].cells[0].text = 'CNN (DL)'
    comparison_table.rows[2].cells[1].text = '95.06%'
    comparison_table.rows[2].cells[2].text = '0.9513'
    
    comparison_table.rows[3].cells[0].text = 'Difference'
    comparison_table.rows[3].cells[1].text = '+3.82%'
    comparison_table.rows[3].cells[2].text = '+0.0377'
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The Logistic Regression model outperformed the CNN model by 3.82% in accuracy. This can be '
        'attributed to the well-structured TF-IDF features and the relatively small dataset size, '
        'where classical ML methods often perform well. The CNN model, while slightly less accurate, '
        'demonstrates the potential for deep learning approaches with larger datasets and more complex patterns.'
    )
    
    # ==================== PART C: TOPIC CLUSTERING ====================
    doc.add_page_break()
    doc.add_heading('5. Part C: Topic Clustering (TF-IDF)', 1)
    
    doc.add_paragraph(
        'Topic clustering was performed to discover latent topics within the news articles using '
        'K-Means clustering on TF-IDF features with dimensionality reduction via SVD.'
    )
    
    doc.add_heading('5.1 Dimensionality Reduction', 2)
    doc.add_paragraph(
        'Truncated SVD (Singular Value Decomposition) was applied to reduce the TF-IDF feature space '
        'from 5,000 dimensions to 100 dimensions for efficient clustering.'
    )
    
    doc.add_paragraph('SVD Configuration:', style='List Bullet')
    doc.add_paragraph('• Original dimensions: 5,000', style='List Bullet 2')
    doc.add_paragraph('• Reduced dimensions: 100', style='List Bullet 2')
    doc.add_paragraph('• Explained variance ratio: 28.39%', style='List Bullet 2')
    
    doc.add_heading('5.2 K-Means Clustering', 2)
    doc.add_paragraph(
        'K-Means clustering was applied with k=5 clusters to match the number of categories. '
        'The algorithm grouped articles into distinct topic clusters.'
    )
    
    cluster_table = doc.add_table(rows=6, cols=3)
    cluster_table.style = 'Light Grid Accent 1'
    cluster_table.rows[0].cells[0].text = 'Cluster ID'
    cluster_table.rows[0].cells[1].text = 'Number of Articles'
    cluster_table.rows[0].cells[2].text = 'Percentage'
    for cell in cluster_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
    
    clusters = [
        ('0', '295', '13.3%'),
        ('1', '379', '17.0%'),
        ('2', '126', '5.7%'),
        ('3', '939', '42.2%'),
        ('4', '486', '21.8%')
    ]
    
    for i, (cid, count, pct) in enumerate(clusters, 1):
        cluster_table.rows[i].cells[0].text = cid
        cluster_table.rows[i].cells[1].text = count
        cluster_table.rows[i].cells[2].text = pct
    
    doc.add_paragraph()
    doc.add_paragraph(
        '[INSERT SCREENSHOT: K-Means Clusters Visualization - PCA projection showing cluster distribution]'
    )
    
    doc.add_heading('5.3 Cluster-Category Analysis', 2)
    doc.add_paragraph(
        'Analysis of cluster-category relationships reveals how well the discovered topics align with '
        'the original categories.'
    )
    
    doc.add_paragraph(
        '[INSERT SCREENSHOT: Cluster-Category Cross-tabulation Table - Showing distribution of categories within each cluster]'
    )
    
    doc.add_heading('5.4 Top Keywords per Cluster', 2)
    
    keywords_table = doc.add_table(rows=6, cols=3)
    keywords_table.style = 'Light Grid Accent 1'
    keywords_table.rows[0].cells[0].text = 'Cluster'
    keywords_table.rows[0].cells[1].text = 'Top Keywords'
    keywords_table.rows[0].cells[2].text = 'Dominant Category'
    for cell in keywords_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
    
    cluster_keywords_data = [
        ('0 (295)', 'mr, labour, election, said, blair, party, government, brown, mr blair, howard', 'Politics (291)'),
        ('1 (379)', 'bn, said, company, shares, mr, firm, market, year, yukos, deal', 'Business (351)'),
        ('2 (126)', 'growth, economy, economic, prices, dollar, said, rate, rates, oil, year', 'Business (126)'),
        ('3 (939)', 'said, film, people, music, new, best, mr, tv, mobile, uk', 'Tech (392)'),
        ('4 (486)', 'game, england, win, said, cup, players, match, play, injury, world', 'Sport (480)')
    ]
    
    for i, (cluster, keywords, category) in enumerate(cluster_keywords_data, 1):
        keywords_table.rows[i].cells[0].text = cluster
        keywords_table.rows[i].cells[1].text = keywords
        keywords_table.rows[i].cells[2].text = category
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The clustering successfully identified distinct topics that largely correspond to the original '
        'categories. Cluster 0 is dominated by politics articles, Cluster 1 and 2 by business articles, '
        'Cluster 3 by tech articles, and Cluster 4 by sport articles.'
    )
    
    # ==================== PART D: REINFORCEMENT LEARNING ====================
    doc.add_page_break()
    doc.add_heading('6. Part D: Reinforcement Learning Decision Agent', 1)
    
    doc.add_paragraph(
        'A Q-Learning reinforcement learning agent was developed to intelligently decide which model '
        'to use (ML or DL) or when to escalate to human review based on the state of the input article.'
    )
    
    doc.add_heading('6.1 State Space Definition', 2)
    doc.add_paragraph(
        'The state space consists of five components that capture relevant information about each article:'
    )
    
    state_components = [
        'ML confidence score: Maximum probability from ML model predictions',
        'DL confidence score: Maximum probability from DL model predictions',
        'Cluster ID: The topic cluster assignment from K-Means',
        'Article length category: Binned article length (short <100, medium 100-200, long >200)',
        'Disagreement flag: Binary indicator if ML and DL predictions differ'
    ]
    
    for i, component in enumerate(state_components, 1):
        doc.add_paragraph(f'{i}. {component}', style='List Number')
    
    doc.add_paragraph()
    doc.add_paragraph('State Space Statistics:', style='List Bullet')
    doc.add_paragraph('• Total states: 750', style='List Bullet 2')
    doc.add_paragraph('• ML confidence range: [0.258, 0.983]', style='List Bullet 2')
    doc.add_paragraph('• DL confidence range: [0.330, 1.000]', style='List Bullet 2')
    doc.add_paragraph('• Disagreement rate: 4.7% (21 out of 445 samples)', style='List Bullet 2')
    
    doc.add_heading('6.2 Action Space', 2)
    doc.add_paragraph('The agent can choose from three actions:')
    
    actions = [
        'Action 0: Use ML prediction',
        'Action 1: Use DL prediction',
        'Action 2: Escalate to human review'
    ]
    
    for action in actions:
        doc.add_paragraph(action, style='List Bullet')
    
    doc.add_heading('6.3 Reward Function', 2)
    doc.add_paragraph('The reward function is designed to encourage correct predictions and appropriate escalation:')
    
    reward_table = doc.add_table(rows=7, cols=2)
    reward_table.style = 'Light Grid Accent 1'
    reward_table.rows[0].cells[0].text = 'Action'
    reward_table.rows[0].cells[1].text = 'Reward'
    reward_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    reward_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    
    rewards = [
        ('ML correct', '+5 + (2 × confidence)'),
        ('ML wrong', '-5'),
        ('DL correct', '+6 + (2 × confidence)'),
        ('DL wrong', '-6'),
        ('Escalate (both wrong)', '+2'),
        ('Escalate (otherwise)', '-1')
    ]
    
    for i, (action, reward) in enumerate(rewards, 1):
        reward_table.rows[i].cells[0].text = action
        reward_table.rows[i].cells[1].text = reward
    
    doc.add_heading('6.4 Q-Learning Training', 2)
    
    doc.add_paragraph('Q-Learning Parameters:', style='List Bullet')
    doc.add_paragraph('• Learning rate (α): 0.1', style='List Bullet 2')
    doc.add_paragraph('• Discount factor (γ): 0.9', style='List Bullet 2')
    doc.add_paragraph('• Exploration rate (ε): 0.2', style='List Bullet 2')
    doc.add_paragraph('• Number of episodes: 1,200', style='List Bullet 2')
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The agent was trained using epsilon-greedy policy, balancing exploration and exploitation. '
        'Training progress showed stable convergence with average rewards around 3,140 per episode.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph(
        '[INSERT SCREENSHOT: Q-Learning Training Progress - Plot showing reward progression over 1,200 episodes]'
    )
    
    doc.add_heading('6.5 RL Agent Evaluation', 2)
    
    doc.add_paragraph('The trained RL agent was evaluated on the test set:')
    
    rl_results_table = doc.add_table(rows=4, cols=2)
    rl_results_table.style = 'Light Grid Accent 1'
    rl_results_table.rows[0].cells[0].text = 'Model/Agent'
    rl_results_table.rows[0].cells[1].text = 'Accuracy'
    rl_results_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    rl_results_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    
    rl_results_table.rows[1].cells[0].text = 'ML Model'
    rl_results_table.rows[1].cells[1].text = '98.88%'
    
    rl_results_table.rows[2].cells[0].text = 'DL Model'
    rl_results_table.rows[2].cells[1].text = '95.06%'
    
    rl_results_table.rows[3].cells[0].text = 'RL Agent'
    rl_results_table.rows[3].cells[1].text = '99.10%'
    
    doc.add_paragraph()
    doc.add_paragraph('Key Findings:', style='List Bullet')
    doc.add_paragraph('• RL Agent achieved the highest accuracy: 99.10%', style='List Bullet 2')
    doc.add_paragraph('• Improvement over ML: +0.22%', style='List Bullet 2')
    doc.add_paragraph('• Improvement over DL: +4.04%', style='List Bullet 2')
    
    doc.add_heading('6.6 Action Distribution', 2)
    doc.add_paragraph('The RL agent\'s decision distribution on the test set:')
    
    action_table = doc.add_table(rows=4, cols=3)
    action_table.style = 'Light Grid Accent 1'
    action_table.rows[0].cells[0].text = 'Action'
    action_table.rows[0].cells[1].text = 'Count'
    action_table.rows[0].cells[2].text = 'Percentage'
    for cell in action_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
    
    action_table.rows[1].cells[0].text = 'Use ML'
    action_table.rows[1].cells[1].text = '18'
    action_table.rows[1].cells[2].text = '4.0%'
    
    action_table.rows[2].cells[0].text = 'Use DL'
    action_table.rows[2].cells[1].text = '427'
    action_table.rows[2].cells[2].text = '96.0%'
    
    action_table.rows[3].cells[0].text = 'Escalate'
    action_table.rows[3].cells[1].text = '0'
    action_table.rows[3].cells[2].text = '0.0%'
    
    doc.add_paragraph()
    doc.add_paragraph(
        '[INSERT SCREENSHOT: RL Action Distribution - Bar chart showing action selection frequency]'
    )
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The agent learned to primarily use the DL model (96% of cases), with occasional use of ML model '
        '(4% of cases). No escalations occurred, suggesting the agent is confident in model predictions.'
    )
    
    doc.add_heading('6.7 Q-Table Statistics', 2)
    doc.add_paragraph('Q-Table Analysis:')
    
    qtable_table = doc.add_table(rows=6, cols=2)
    qtable_table.style = 'Light Grid Accent 1'
    qtable_table.rows[0].cells[0].text = 'Metric'
    qtable_table.rows[0].cells[1].text = 'Value'
    qtable_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    qtable_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    
    qtable_stats = [
        ('Q-table shape', '750 × 3'),
        ('Min Q-value', '0.0000'),
        ('Max Q-value', '79.9995'),
        ('Mean Q-value', '4.7046'),
        ('Non-zero states', '52 / 750 (6.9%)')
    ]
    
    for i, (metric, value) in enumerate(qtable_stats, 1):
        qtable_table.rows[i].cells[0].text = metric
        qtable_table.rows[i].cells[1].text = value
    
    # ==================== RESULTS AND ANALYSIS ====================
    doc.add_page_break()
    doc.add_heading('7. Results and Analysis', 1)
    
    doc.add_heading('7.1 Overall Performance Summary', 2)
    
    summary_table = doc.add_table(rows=4, cols=4)
    summary_table.style = 'Light Grid Accent 1'
    summary_table.rows[0].cells[0].text = 'Model'
    summary_table.rows[0].cells[1].text = 'Accuracy'
    summary_table.rows[0].cells[2].text = 'F1-Score (Macro)'
    summary_table.rows[0].cells[3].text = 'Misclassified'
    for cell in summary_table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
    
    summary_table.rows[1].cells[0].text = 'Logistic Regression'
    summary_table.rows[1].cells[1].text = '98.88%'
    summary_table.rows[1].cells[2].text = '0.9890'
    summary_table.rows[1].cells[3].text = '5'
    
    summary_table.rows[2].cells[0].text = 'CNN (Deep Learning)'
    summary_table.rows[2].cells[1].text = '95.06%'
    summary_table.rows[2].cells[2].text = '0.9513'
    summary_table.rows[2].cells[3].text = '22'
    
    summary_table.rows[3].cells[0].text = 'RL Agent'
    summary_table.rows[3].cells[1].text = '99.10%'
    summary_table.rows[3].cells[2].text = 'N/A'
    summary_table.rows[3].cells[3].text = '4'
    
    doc.add_heading('7.2 Key Insights', 2)
    
    insights = [
        'The Logistic Regression model achieved excellent performance (98.88% accuracy), demonstrating '
        'the effectiveness of TF-IDF features for text classification.',
        
        'The CNN model, while slightly less accurate, shows promise for handling more complex patterns '
        'and could benefit from larger datasets or more sophisticated architectures.',
        
        'The Reinforcement Learning agent successfully learned to combine both models, achieving the '
        'highest accuracy (99.10%) by intelligently selecting the best model for each article.',
        
        'Topic clustering revealed distinct themes that largely align with the original categories, '
        'with some clusters showing clear dominance of specific categories.',
        
        'The RL agent primarily favored the DL model (96% of cases), suggesting it learned that the '
        'DL model\'s predictions were more reliable in most scenarios, despite the ML model\'s higher '
        'overall accuracy.'
    ]
    
    for i, insight in enumerate(insights, 1):
        doc.add_paragraph(f'{i}. {insight}')
    
    doc.add_heading('7.3 Limitations and Future Work', 2)
    
    limitations = [
        'The dataset size (2,225 articles) is relatively small for deep learning, which may limit '
        'the CNN model\'s potential.',
        
        'The RL agent did not escalate any cases, suggesting the reward function may need tuning to '
        'encourage more conservative escalation when confidence is low.',
        
        'Topic clustering could be improved with more sophisticated techniques like LDA or BERT-based '
        'embeddings for better semantic understanding.',
        
        'Hyperparameter tuning could further improve model performance, especially for the CNN architecture.',
        
        'The state space encoding could be refined to capture more nuanced information about article '
        'characteristics.'
    ]
    
    for limitation in limitations:
        doc.add_paragraph(limitation, style='List Bullet')
    
    # ==================== CONCLUSION ====================
    doc.add_page_break()
    doc.add_heading('8. Conclusion', 1)
    
    doc.add_paragraph(
        'This project successfully implemented a comprehensive news classification system combining '
        'classical machine learning, deep learning, and reinforcement learning techniques. The system '
        'demonstrates the effectiveness of hybrid approaches, where the RL agent intelligently combines '
        'the strengths of both ML and DL models to achieve superior performance.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Key achievements include:'
    )
    
    achievements = [
        'Successfully classified news articles with 99.10% accuracy using the RL agent',
        'Identified distinct topic clusters that align with article categories',
        'Demonstrated the value of reinforcement learning in model selection',
        'Provided a framework for intelligent decision-making in classification tasks'
    ]
    
    for achievement in achievements:
        doc.add_paragraph(achievement, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The project showcases the practical application of multiple machine learning paradigms and '
        'demonstrates how they can be combined to create more robust and intelligent classification systems.'
    )
    
    # ==================== REFERENCES ====================
    doc.add_page_break()
    doc.add_heading('9. References', 1)
    
    references = [
        'BBC News Dataset. (n.d.). Kaggle. https://www.kaggle.com/datasets/pariza/bbc-news-summary',
        'Scikit-learn: Machine Learning in Python. Pedregosa et al., JMLR 12, pp. 2825-2830, 2011.',
        'TensorFlow: Large-Scale Machine Learning on Heterogeneous Systems. Abadi et al., 2015.',
        'Sutton, R. S., & Barto, A. G. (2018). Reinforcement Learning: An Introduction. MIT Press.',
        'Manning, C. D., Raghavan, P., & Schütze, H. (2008). Introduction to Information Retrieval. Cambridge University Press.'
    ]
    
    for i, ref in enumerate(references, 1):
        doc.add_paragraph(f'[{i}] {ref}')
    
    # Save document
    output_path = os.path.join(os.path.dirname(__file__), 'News_Classification_Report.docx')
    doc.save(output_path)
    print(f"Report generated successfully: {output_path}")
    return output_path

if __name__ == '__main__':
    create_report()

